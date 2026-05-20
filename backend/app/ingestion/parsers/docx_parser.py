import io
import logging
from typing import Dict, Any, List, Tuple

logger = logging.getLogger(__name__)

try:
    import docx
except ImportError:
    docx = None


def extract_docx_content(file_bytes: bytes, filename: str) -> Tuple[str, List[Dict[str, Any]], Dict[str, Any]]:
    """
    Parses a DOCX document, retaining headings, paragraphs, and tables.
    """
    full_text_list = []
    structural_meta = []
    global_meta = {"filename": filename}

    if not docx:
        logger.error("python-docx is not installed.")
        return file_bytes.decode("utf-8", errors="ignore"), [], global_meta

    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        
        current_section = "Introduction"
        element_idx = 0
        
        # Iterate over structural elements
        for block in doc.element.body:
            # Check if block is a paragraph
            if block.tag.endswith('p'):
                p = docx.text.paragraph.Paragraph(block, doc)
                p_text = p.text.strip()
                if not p_text:
                    continue
                
                # Check style for headings
                style_name = p.style.name if p.style else ""
                is_heading = style_name.startswith("Heading") or p.style.name.startswith("Title")
                
                if is_heading:
                    current_section = p_text
                    full_text_list.append(f"\n# {p_text}\n")
                else:
                    full_text_list.append(p_text)
                
                structural_meta.append({
                    "element_index": element_idx,
                    "type": "heading" if is_heading else "paragraph",
                    "section_title": current_section,
                    "char_count": len(p_text)
                })
                element_idx += 1
                
            # Check if block is a table
            elif block.tag.endswith('tbl'):
                # Extract docx.table.Table objects
                table_idx = len([m for m in structural_meta if m["type"] == "table"])
                if table_idx < len(doc.tables):
                    table = doc.tables[table_idx]
                    table_rows = []
                    for row in table.rows:
                        row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                        if any(row_cells):
                            table_rows.append(" | ".join(row_cells))
                    
                    if table_rows:
                        table_text = "\n[Table Extraction]\n" + "\n".join(table_rows) + "\n"
                        full_text_list.append(table_text)
                        
                        structural_meta.append({
                            "element_index": element_idx,
                            "type": "table",
                            "section_title": current_section,
                            "char_count": len(table_text)
                        })
                        element_idx += 1

    except Exception as e:
        logger.error(f"Error parsing DOCX {filename}: {e}")
        # Basic raw text extraction fallback
        return file_bytes.decode("utf-8", errors="ignore"), [], global_meta

    full_text = "\n\n".join(full_text_list)
    return full_text, structural_meta, global_meta
